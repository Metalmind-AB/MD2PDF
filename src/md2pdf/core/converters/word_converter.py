#!/usr/bin/env python3
"""
MD2PDF - Markdown to PDF Converter
Copyright (c) 2025 MPS Metalmind AB
Licensed under the MIT License (see LICENSE file)
"""

"""
Word Converter - Converts Markdown to Word using python-docx
Inherits from BaseConverter and provides Word-specific functionality.
"""

from pathlib import Path

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.text.paragraph import Paragraph

from md2pdf.core.converters.base_converter import BaseConverter


class WordConverter(BaseConverter):
    """Word converter using python-docx."""

    def _generate_output_path(self) -> Path:
        """Generate output Word path based on input file."""
        return self.input_file.with_suffix(".docx")

    def _ensure_docx_extension(self, output_path: Path) -> Path:
        """Ensure the output path has .docx extension."""
        if output_path.suffix.lower() != ".docx":
            return output_path.with_suffix(".docx")
        return output_path

    def _apply_style_to_paragraph(
        self, paragraph: Paragraph, element: Tag, style_name: str
    ) -> None:
        """Apply style to paragraph based on HTML element."""
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            if level == 1:
                paragraph.style = "Heading 1"
            elif level == 2:
                paragraph.style = "Heading 2"
            elif level == 3:
                paragraph.style = "Heading 3"
            else:
                paragraph.style = f"Heading {level}"
        elif element.name == "p":
            paragraph.style = "Normal"
        elif element.name == "strong":
            paragraph.style = "Strong"
        elif element.name == "em":
            paragraph.style = "Emphasis"

    def _convert_html_to_word(self, html_content: str) -> Document:
        """Convert HTML content to Word document."""
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        # Remove script and style elements
        self._clean_soup(soup)

        # Process content
        self._process_elements(doc, soup)

        return doc

    def _clean_soup(self, soup: BeautifulSoup) -> None:
        """Remove script and style elements from soup."""
        for script in soup(["script", "style"]):
            script.decompose()

    def _process_elements(self, doc: Document, soup: BeautifulSoup) -> None:
        """Process HTML elements and add to Word document."""
        for element in soup.find_all(
            [
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "p",
                "ul",
                "ol",
                "li",
                "blockquote",
                "pre",
                "code",
            ]
        ):
            if not isinstance(element, Tag) or not element.name:
                continue

            self._process_single_element(doc, element)

    def _process_single_element(self, doc: Document, element: Tag) -> None:
        """Process a single HTML element."""
        element_name = element.name

        if element_name in ["h1", "h2", "h3", "h4", "h5", "h6", "p"]:
            self._add_text_paragraph(doc, element)
        elif element_name in ["ul", "ol"]:
            self._add_list_elements(doc, element)
        elif element_name == "blockquote":
            self._add_blockquote(doc, element)
        elif element_name in ["pre", "code"]:
            self._add_code_block(doc, element)

    def _add_text_paragraph(self, doc: Document, element: Tag) -> None:
        """Add text paragraph (headings, paragraphs) to document."""
        paragraph = doc.add_paragraph()
        paragraph.text = element.get_text()
        self._apply_style_to_paragraph(paragraph, element, element.name)

    def _add_list_elements(self, doc: Document, element: Tag) -> None:
        """Add list elements to document."""
        list_style = "List Bullet" if element.name == "ul" else "List Number"

        for li in element.find_all("li", recursive=False):
            if isinstance(li, Tag):
                paragraph = doc.add_paragraph()
                paragraph.text = li.get_text()
                paragraph.style = list_style

    def _add_blockquote(self, doc: Document, element: Tag) -> None:
        """Add blockquote element to document."""
        paragraph = doc.add_paragraph()
        paragraph.text = element.get_text()
        paragraph.style = "Quote"

    def _add_code_block(self, doc: Document, element: Tag) -> None:
        """Add code block element to document."""
        paragraph = doc.add_paragraph()
        paragraph.text = element.get_text()
        paragraph.style = "No Spacing"
        # Apply monospace font
        for run in paragraph.runs:
            run.font.name = "Courier New"

    def convert(self) -> bool:
        """Convert Markdown file to Word document."""
        try:
            # Ensure output path has .docx extension
            self.output_file: Path = self._ensure_docx_extension(self.output_file)

            # Read the markdown file
            markdown_content = self._read_markdown_content()

            # Convert markdown to HTML
            html_content = self._process_markdown(markdown_content)

            # Convert HTML to Word
            print(f"Generating Word document: {self.output_file}")
            doc = self._convert_html_to_word(html_content)

            # Save the document
            doc.save(str(self.output_file))

            print(f"✅ Successfully created Word document: {self.output_file}")
            return True

        except Exception as e:
            print(f"❌ Error converting to Word: {str(e)}")
            return False
