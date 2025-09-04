#!/usr/bin/env python3
"""
MD2PDF - Markdown to PDF Converter
Copyright (c) 2025 MPS Metalmind AB
Licensed under the MIT License (see LICENSE file)
"""

"""
Word Converter - Converts Markdown to Word using python-docx
Inherits from BaseConverter and provides Word-specific functionality.
"""

from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document

from .base_converter import BaseConverter


class WordConverter(BaseConverter):
    """Word converter using python-docx."""

    def _generate_output_path(self) -> Path:
        """Generate output Word path based on input file."""
        return self.input_file.with_suffix(".docx")

    def _ensure_docx_extension(self, output_path: Path) -> Path:
        """Ensure the output path has .docx extension."""
        if output_path.suffix.lower() != ".docx":
            return output_path.with_suffix(".docx")
        return output_path

    def _apply_style_to_paragraph(self, paragraph, element, style_name):
        """Apply style to paragraph based on HTML element."""
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            if level == 1:
                paragraph.style = "Heading 1"
            elif level == 2:
                paragraph.style = "Heading 2"
            elif level == 3:
                paragraph.style = "Heading 3"
            else:
                paragraph.style = f"Heading {level}"
        elif element.name == "p":
            paragraph.style = "Normal"
        elif element.name == "strong":
            paragraph.style = "Strong"
        elif element.name == "em":
            paragraph.style = "Emphasis"

    def _convert_html_to_word(self, html_content: str) -> Document:
        """Convert HTML content to Word document."""
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Process content
        for element in soup.find_all(
            [
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "p",
                "ul",
                "ol",
                "li",
                "blockquote",
                "pre",
                "code",
            ]
        ):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6", "p"]:
                paragraph = doc.add_paragraph()
                paragraph.text = element.get_text()
                self._apply_style_to_paragraph(
                    paragraph, element, element.name
                )

            elif element.name in ["ul", "ol"]:
                for li in element.find_all("li", recursive=False):
                    paragraph = doc.add_paragraph()
                    paragraph.text = li.get_text()
                    if element.name == "ul":
                        paragraph.style = "List Bullet"
                    else:
                        paragraph.style = "List Number"

            elif element.name == "blockquote":
                paragraph = doc.add_paragraph()
                paragraph.text = element.get_text()
                paragraph.style = "Quote"

            elif element.name in ["pre", "code"]:
                paragraph = doc.add_paragraph()
                paragraph.text = element.get_text()
                paragraph.style = "No Spacing"
                # Apply monospace font
                for run in paragraph.runs:
                    run.font.name = "Courier New"

        return doc

    def convert(self) -> bool:
        """Convert Markdown file to Word document."""
        try:
            # Ensure output path has .docx extension
            self.output_file = self._ensure_docx_extension(self.output_file)

            # Read the markdown file
            markdown_content = self._read_markdown_content()

            # Convert markdown to HTML
            html_content = self._process_markdown(markdown_content)

            # Convert HTML to Word
            print(f"Generating Word document: {self.output_file}")
            doc = self._convert_html_to_word(html_content)

            # Save the document
            doc.save(str(self.output_file))

            print(f"✅ Successfully created Word document: {self.output_file}")
            return True

        except Exception as e:
            print(f"❌ Error converting to Word: {str(e)}")
            return False
